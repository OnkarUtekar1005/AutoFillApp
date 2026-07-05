"""Extract line items from Word (.docx) financial documents."""
import io
import re

import docx


def _is_numeric_cell(value: str) -> bool:
    cleaned = re.sub(r'[₹$,\s()\-]', '', str(value))
    try:
        float(cleaned)
        return bool(cleaned)
    except ValueError:
        return False


def _is_short_generic(label: str) -> bool:
    s = label.strip()
    if len(s) <= 2:
        return True
    if re.fullmatch(r'[IVXivx]+\.?', s):
        return True
    if re.fullmatch(r'\([a-zA-Z]\)', s):
        return True
    return False


def extract_word(file_bytes: bytes) -> list[dict]:
    """
    Extract label-value pairs from a Word document.
    Processes tables first (most financial data), then paragraphs.
    Captures both current-year and prior-year columns from tables.
    """
    items: list[dict] = []
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text_parts: list[str] = []

    # ── Tables ────────────────────────────────────────────────────────────────
    for table in doc.tables:
        row_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not row_data:
            continue

        # Detect label column
        col_text_scores: dict[int, int] = {}
        for row in row_data:
            for ci, cell in enumerate(row):
                if cell and not _is_numeric_cell(cell) and len(cell) > 3:
                    col_text_scores[ci] = col_text_scores.get(ci, 0) + 1

        if not col_text_scores:
            continue

        label_col = max(col_text_scores, key=col_text_scores.get)

        # Detect CY / PY columns from header row
        cy_col: int | None = None
        py_col: int | None = None
        if row_data:
            header = row_data[0]
            year_cols = [
                ci for ci, c in enumerate(header)
                if ci > label_col and re.search(r'20\d{2}', c)
            ]
            if len(year_cols) >= 2:
                cy_col, py_col = year_cols[0], year_cols[1]
            elif len(year_cols) == 1:
                cy_col = year_cols[0]

        for row in row_data:
            if label_col >= len(row):
                continue
            label = row[label_col]
            if not label or _is_numeric_cell(label) or _is_short_generic(label):
                continue

            cy_val = ""
            if cy_col is not None and cy_col < len(row) and _is_numeric_cell(row[cy_col]):
                cy_val = row[cy_col]
            if not cy_val:
                for ci in range(label_col + 1, min(label_col + 6, len(row))):
                    if _is_numeric_cell(row[ci]):
                        cy_val = row[ci]
                        if cy_col is None:
                            cy_col = ci
                        break

            py_val = ""
            if py_col is not None and py_col < len(row) and _is_numeric_cell(row[py_col]):
                py_val = row[py_col]
            if not py_val and cy_col is not None:
                for ci in range(cy_col + 1, min(cy_col + 4, len(row))):
                    if ci < len(row) and _is_numeric_cell(row[ci]):
                        py_val = row[ci]
                        break

            if cy_val:
                items.append({"raw_label": label, "raw_value": cy_val,
                               "page": None, "bbox": None, "ocr_confidence": None})
            if py_val:
                items.append({"raw_label": "previous year " + label, "raw_value": py_val,
                               "page": None, "bbox": None, "ocr_confidence": None})

    # ── Paragraphs ────────────────────────────────────────────────────────────
    line_pattern = re.compile(r'^(.{4,60}?)\s{2,}([\d,₹$\(\)\-\.]+)\s*$')
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        full_text_parts.append(text)
        m = line_pattern.match(text)
        if m:
            label, value = m.group(1).strip(), m.group(2).strip()
            if _is_numeric_cell(value) and not _is_short_generic(label):
                items.append({"raw_label": label, "raw_value": value,
                               "page": None, "bbox": None, "ocr_confidence": None})

    # Non-numeric fields from full document text
    from extractor.extraction.text_fields import extract_text_fields
    items.extend(extract_text_fields("\n".join(full_text_parts)))
    return items
